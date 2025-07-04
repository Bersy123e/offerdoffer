# django_app/products/client_request_extractor.py

import pandas as pd
import re
import logging
import numpy as np
from typing import List, Dict, Optional, Any, Tuple
import os
import json
import csv
import docx
import tabula # Для извлечения таблиц из PDF
import openpyxl # Для детального анализа стилей
import xlrd # Для чтения старых файлов .xls
from pydantic import BaseModel, ValidationError, Field
from langchain.prompts import PromptTemplate # Добавил импорт PromptTemplate
from langchain.chains import LLMChain # Добавил импорт LLMChain

# Настройка логгера для нового модуля
logger = logging.getLogger('commercial_proposal')

# --- Pydantic модели для валидации извлеченных товаров клиента ---

class ClientRequestedItem(BaseModel):
    """Модель для валидации одного товара, извлеченного из запроса клиента."""
    full_name: str = Field(..., min_length=3)
    quantity: int = Field(..., ge=0) # Количество должно быть неотрицательным

# --- Основной класс для извлечения данных из запросов клиентов ---

class ClientRequestExtractor:
    """
    Каскадный процессор для извлечения товаров и количества из запросов клиентов.
    Адаптирован на основе CascadeProcessor для прайс-листов.
    """

    def __init__(self, llm: Any):
        if not llm:
            raise ValueError("LLM instance is required.")
        self.llm = llm
        self.cascade_log = []

    def process_client_request_file_cascade(self, file_path: str, file_name: str) -> Dict:
        """
        Каскадная обработка файла запроса клиента с тремя уровнями.
        
        УРОВЕНЬ 1: Пространственный анализ LLM (для сложных форматов)
        УРОВЕНЬ 2: Структурный анализ LLM + извлечение (для стандартных таблиц)
        УРОВЕНЬ 3: Эвристические методы (fallback для простых случаев)
        """
        logger.info(f"--- ЗАПУСК КАСКАДНОЙ СИСТЕМЫ для клиентского запроса: {file_name} ---")
        self.cascade_log = [f"Начало каскадной обработки запроса: {file_name}"]
        
        try:
            ext = os.path.splitext(file_path)[1].lower()
            self.cascade_log.append(f"Определен формат файла: {ext}")
            
            # === УРОВЕНЬ 1: ПРОСТРАНСТВЕННЫЙ АНАЛИЗ LLM ===
            # Этот уровень подходит для всех форматов, которые могут содержать таблицы
            if ext in ['.xlsx', '.xls', '.pdf', '.docx']: # Добавили PDF и DOCX
                self.cascade_log.append("УРОВЕНЬ 1: Пространственный анализ LLM для запроса")
                level1_result = self._process_level1_spatial_client(file_path, file_name)
                
                if level1_result["success"] and len(level1_result["items"]) >= 1: # Минимум 1 товар
                    self.cascade_log.append(f"УРОВЕНЬ 1 УСПЕШЕН: {len(level1_result['items'])} позиций")
                    level1_result["final_method"] = "Level 1: LLM Spatial Analysis (Client Request)"
                    level1_result["cascade_log"] = self.cascade_log
                    return level1_result
                else:
                    self.cascade_log.append(f"УРОВЕНЬ 1: {len(level1_result.get('items', []))} позиций (недостаточно или ошибка)")
            else:
                level1_result = {"success": False, "items": [], "error": f"Формат {ext} не поддерживается Уровнем 1"}


            # === УРОВЕНЬ 2: СТРУКТУРНЫЙ АНАЛИЗ LLM ===
            self.cascade_log.append("УРОВЕНЬ 2: Структурный анализ LLM для запроса")
            level2_result = self._process_level2_structural_client(file_path, file_name)
            
            if level2_result["success"] and len(level2_result["items"]) >= 1: # Минимум 1 товар
                self.cascade_log.append(f"УРОВЕНЬ 2 УСПЕШЕН: {len(level2_result['items'])} позиций")
                level2_result["final_method"] = "Level 2: LLM Structural Analysis (Client Request)"
                level2_result["cascade_log"] = self.cascade_log
                return level2_result
            else:
                self.cascade_log.append(f"УРОВЕНЬ 2: {len(level2_result.get('items', []))} позиций")
            
            # === УРОВЕНЬ 3: ЭВРИСТИЧЕСКИЕ МЕТОДЫ ===
            self.cascade_log.append("УРОВЕНЬ 3: Эвристические методы (fallback) для запроса")
            level3_result = self._process_level3_heuristics_client(file_path, file_name)
            
            if level3_result["success"] and len(level3_result["items"]) > 0:
                self.cascade_log.append(f"УРОВЕНЬ 3 УСПЕШЕН: {len(level3_result['items'])} позиций")
                level3_result["final_method"] = "Level 3: Heuristic Methods (Client Request)"
                level3_result["cascade_log"] = self.cascade_log
                return level3_result
            else:
                self.cascade_log.append(f"УРОВЕНЬ 3: {len(level3_result.get('items', []))} позиций")
            
            # === ВЫБОР ЛУЧШЕГО РЕЗУЛЬТАТА ===
            results = []
            # Убеждаемся, что level1_result был создан, если ext был в списке
            if level1_result.get("success"): results.append(("Level 1", level1_result))
            if level2_result.get("success"): results.append(("Level 2", level2_result))
            if level3_result.get("success"): results.append(("Level 3", level3_result))
            
            if results:
                best_name, best_result = max(results, key=lambda x: len(x[1].get("items", [])))
                self.cascade_log.append(f"ВЫБРАН ЛУЧШИЙ РЕЗУЛЬТАТ: {best_name} с {len(best_result['items'])} позициями")
                best_result["final_method"] = f"Best of All: {best_name} (Client Request)"
                best_result["cascade_log"] = self.cascade_log
                return best_result
            
            # Если все методы провалились
            error_msg = "Все уровни каскадной обработки не смогли извлечь данные из клиентского запроса."
            return self._handle_error(error_msg, is_client_request=True) # Добавим флаг для логгера
            
        except Exception as e:
            return self._handle_error(f"Критическая ошибка в каскадной системе для клиентского запроса: {e}", exc_info=True, is_client_request=True)

    def _process_level1_spatial_client(self, file_path: str, file_name: str) -> Dict:
        """УРОВЕНЬ 1 для клиентского запроса: Пространственный анализ LLM."""
        try:
            spatial_json = self._file_to_spatial_json(file_path)
            if not spatial_json:
                return {"success": False, "items": [], "error": "Не удалось создать пространственное представление"}

            client_items_json = self._get_client_items_from_llm(spatial_json)
            if not client_items_json:
                return {"success": False, "items": [], "error": "LLM не смогла извлечь данные из пространственного JSON"}

            validated_items = self._validate_and_clean_client_items(client_items_json, "Level 1 (Client)")
            
            if not validated_items:
                return {"success": False, "items": [], "error": "Ни одна позиция не прошла валидацию"}

            return {"success": True, "items": validated_items}
            
        except Exception as e:
            return {"success": False, "items": [], "error": f"Ошибка уровня 1 (клиент): {e}"}

    def _process_level2_structural_client(self, file_path: str, file_name: str) -> Dict:
        """УРОВЕНЬ 2 для клиентского запроса: Структурный анализ LLM."""
        try:
            # Используем универсальный _file_to_dataframe для чтения Excel/CSV
            df = self._file_to_dataframe(file_path)
            if df is None:
                self.cascade_log.append(f"УРОВЕНЬ 2: Не удалось прочитать файл {file_name} в DataFrame.")
                return {"success": False, "items": [], "error": "Не удалось прочитать файл"}

            # Для запросов клиента мы можем использовать упрощенный структурный анализ
            # Промпт для LLM должен быть адаптирован для поиска имени и количества
            
            # Если файл "большой", берем репрезентативную выборку.
            if len(df) > 200:
                df_non_empty = df.dropna(how='all')
                sample_df = pd.concat([
                    df_non_empty.head(30), 
                    df_non_empty.sample(n=min(30, len(df_non_empty)), random_state=1), 
                    df_non_empty.tail(30)
                ]).drop_duplicates()
                sample_text = sample_df.to_csv(index=False, header=False, lineterminator='\n')
            else:
                sample_text = df.to_csv(index=False, header=False, lineterminator='\n')

            client_items_json = self._get_client_items_from_text(sample_text)
            if not client_items_json:
                return {"success": False, "items": [], "error": "LLM не смогла извлечь данные из текста"}

            validated_items = self._validate_and_clean_client_items(client_items_json, "Level 2 (Client)")
            
            if not validated_items:
                return {"success": False, "items": [], "error": "Ни одна позиция не прошла валидацию"}

            return {"success": True, "items": validated_items}
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка уровня 2 (клиент): {e}")
            return {"success": False, "items": [], "error": f"Ошибка уровня 2 (клиент): {e}"}

    def _process_level3_heuristics_client(self, file_path: str, file_name: str) -> Dict:
        """УРОВЕНЬ 3 для клиентского запроса: Эвристические методы (fallback)."""
        try:
            extracted_text = ""
            ext = os.path.splitext(file_path)[1].lower()
            
            # Унифицированное чтение текста из разных форматов для эвристик
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
            elif ext == '.docx':
                doc = docx.Document(file_path)
                extracted_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()] +
                                            [cell.text for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()])
            elif ext == '.pdf':
                try:
                    # Попытка извлечь таблицы с tabula-py, затем текст с PyPDF2
                    tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True, pandas_options={'header': None})
                    if tables:
                        extracted_text = '\n'.join([df.to_string(index=False, header=False) for df in tables])
                    else:
                        from PyPDF2 import PdfReader
                        reader = PdfReader(file_path)
                        extracted_text = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])
                except Exception as pdf_err:
                    self.cascade_log.append(f"Ошибка чтения PDF (Tabula/PyPDF2): {pdf_err}")
                    return {"success": False, "items": [], "error": f"Ошибка чтения PDF: {pdf_err}"}
            elif ext in ['.xlsx', '.xls', '.csv']:
                # Используем универсальный _file_to_dataframe для чтения Excel/CSV
                df = self._file_to_dataframe(file_path)
                if df is not None:
                    extracted_text = df.to_string(index=False, header=False) # Преобразуем DataFrame в строку
                else:
                    self.cascade_log.append(f"УРОВЕНЬ 3: Не удалось прочитать файл {file_name} в DataFrame для эвристического анализа.")
                    return {"success": False, "items": [], "error": f"Не удалось прочитать файл для эвристического анализа: {file_name}"}
            else:
                self.cascade_log.append(f"Формат {ext} не поддерживается для эвристического анализа.")
                return {"success": False, "items": [], "error": f"Неподдерживаемый формат для эвристик: {ext}"}
            
            if not extracted_text.strip():
                return {"success": False, "items": [], "error": "Не удалось извлечь текст для эвристического анализа"}

            items = []
            # Простые эвристики для поиска "Название - Количество"
            # Ищем строки, содержащие потенциальное название и число
            lines = extracted_text.split('\n')
            for line in lines:
                line = line.strip()
                if len(line) < 5 or "цена" in line.lower() or "руб" in line.lower(): # Пропускаем слишком короткие или строки с ценой
                    continue
                
                # Ищем число (количество) в конце строки
                qty_match = re.search(r'(\d+)\s*(?:шт|штук|компл|ед|м|тонн|\b)\s*$', line, re.IGNORECASE)
                if qty_match:
                    try:
                        quantity = int(qty_match.group(1))
                        name = re.sub(r'(\d+)\s*(?:шт|штук|компл|ед|м|тонн|\b)\s*$', '', line, flags=re.IGNORECASE).strip()
                        if name and len(name) >= 3:
                            items.append({"full_name": name, "quantity": quantity})
                    except ValueError:
                        pass
                else:
                    # Если явного количества нет, ищем просто последнее число в строке
                    numbers = re.findall(r'\d+', line)
                    if numbers:
                        last_number = int(numbers[-1])
                        # Простая эвристика: если число не слишком большое (не похоже на год или артикул)
                        if last_number < 10000 and last_number > 0: # Ограничим количество, чтобы избежать ложных срабатываний
                            name_part = re.sub(r'\d+', '', line).strip() # Убираем все числа
                            if name_part and len(name_part) >= 3:
                                items.append({"full_name": name_part, "quantity": last_number})

            validated_items = self._validate_and_clean_client_items(items, "Level 3 (Client)")
            
            if not validated_items:
                return {"success": False, "items": [], "error": "Ни одна позиция не прошла валидацию в эвристиках"}

            return {"success": True, "items": validated_items}
                
        except Exception as e:
            self.cascade_log.append(f"Ошибка уровня 3 (клиент): {e}")
            return {"success": False, "items": [], "error": f"Ошибка уровня 3 (клиент): {e}"}

    def _validate_and_clean_client_items(self, items: List[Dict], level_name: str) -> List[Dict]:
        """
        Универсальная валидация и очистка извлеченных позиций клиента.
        
        Выполняет:
        1. Pydantic валидацию
        2. Проверку качества данных
        3. Очистку и нормализацию
        4. Фильтрацию дубликатов
        """
        validated_items = []
        skipped_count = 0
        
        for i, item in enumerate(items):
            try:
                # Нормализация структуры данных (если нужно)
                normalized_item = {
                    "full_name": str(item.get("full_name", item.get("name", ""))).strip(),
                    "quantity": int(item.get("quantity", 0)) # Количество приводим к int
                }
                
                # Pydantic валидация
                validated_item = ClientRequestedItem(**normalized_item)
                
                # Проверка качества данных (адаптирована под клиентский запрос)
                if self._is_quality_client_item(validated_item):
                    validated_items.append(validated_item.dict())
                else:
                    skipped_count += 1
                    self.cascade_log.append(f"{level_name}: Пропущена позиция низкого качества: {normalized_item.get('full_name', 'Без названия')[:50]}")
                    
            except (ValidationError, ValueError, TypeError) as e:
                skipped_count += 1
                self.cascade_log.append(f"{level_name}: Ошибка валидации позиции {i+1}: {e} (Данные: {item})")
                continue
            except Exception as e:
                skipped_count += 1
                self.cascade_log.append(f"{level_name}: Неожиданная ошибка при валидации позиции {i+1}: {e}")
                continue
        
        # Удаление дубликатов (адаптировано для названия + количества)
        unique_items = self._remove_client_item_duplicates(validated_items)
        
        removed_duplicates = len(validated_items) - len(unique_items)
        if removed_duplicates > 0:
            self.cascade_log.append(f"{level_name}: Удалено {removed_duplicates} дубликатов позиций")
        
        self.cascade_log.append(f"{level_name}: Валидация завершена. Принято: {len(unique_items)}, Отклонено: {skipped_count}")
        
        return unique_items

    def _is_quality_client_item(self, item: ClientRequestedItem) -> bool:
        """
        Проверяет качество позиции клиента по множественным критериям.
        
        Критерии качества:
        1. Название не менее 3 символов
        2. Название не является служебным словом
        3. Количество > 0 и разумно (например, до 1000000)
        4. Название содержит значимую информацию
        """
        name = item.full_name.lower().strip()
        
        if len(name) < 3:
            return False
        
        service_words = [
            'nan', 'none', 'null', 'undefined', 'наименование', 'товар', 'продукт',
            'название', 'описание', 'итого', 'всего', 'сумма', 'total', 'sum',
            'заголовок', 'header', 'title', 'примечание', 'note', 'комментарий',
            'список', 'позиция', 'артикул', 'счет'
        ]
        
        if name in service_words or (len(name.split()) < 2 and not any(char.isdigit() for char in name)):
            return False
        
        # Проверка количества
        if not (item.quantity > 0 and item.quantity <= 1000000): # Разумные пределы
            return False
        
        return True

    def _remove_client_item_duplicates(self, items: List[Dict]) -> List[Dict]:
        """Удаляет ИСТИННЫЕ дубликаты позиций клиента (одинаковые название + количество)."""
        seen_combinations = set()
        unique_items = []
        
        for item in items:
            name_key = item["full_name"].lower().strip()
            qty_key = str(item.get("quantity", ""))
            
            combination_key = f"{name_key}|{qty_key}"
            
            if combination_key not in seen_combinations:
                seen_combinations.add(combination_key)
                unique_items.append(item)
        
        return unique_items

    def _handle_error(self, message: str, exc_info=False, is_client_request: bool = False) -> Dict:
        """Централизованный обработчик ошибок."""
        logger.error(message, exc_info=exc_info)
        log_prefix = "ОШИБКА (КЛИЕНТ):" if is_client_request else "ОШИБКА:"
        self.cascade_log.append(f"{log_prefix} {message}")
        # Возвращаем items вместо products
        return {"success": False, "items": [], "error": message, "cascade_log": self.cascade_log}

    # --- Копируем универсальные методы чтения файлов (без изменений) ---
    def _file_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует файл любого поддерживаемого формата в детализированный JSON с координатами и стилями."""
        self.cascade_log.append("Шаг 1: Создание пространственного JSON представления.")
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext in ['.xlsx', '.xls']:
                return self._excel_to_spatial_json(file_path)
            
            elif ext == '.csv':
                return self._csv_to_spatial_json(file_path)
            
            elif ext == '.pdf':
                return self._pdf_to_spatial_json(file_path)
            
            elif ext == '.docx':
                return self._docx_to_spatial_json(file_path)
            
            else:
                self.cascade_log.append(f"Формат {ext} не поддерживается для пространственного анализа.")
                return None
                
        except Exception as e:
            self.cascade_log.append(f"Ошибка при создании пространственного JSON: {e}")
            return None

    def _excel_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует Excel файл в пространственный JSON."""
        try:
            # openpyxl не поддерживает .xls, поэтому здесь только .xlsx
            if file_path.lower().endswith('.xls'):
                self.cascade_log.append(f"Ошибка при обработке Excel файла: openpyxl does not support the old .xls file format, please use xlrd to read this file, or convert it to the more recent .xlsx file format.")
                return None
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            sheet = workbook.active
            
            cells_data = []
            max_rows_to_process = 500  # Достаточно для большинства запросов
            for row_idx, row in enumerate(sheet.iter_rows(max_row=max_rows_to_process)):
                for col_idx, cell in enumerate(row):
                    if cell.value is not None:
                        cells_data.append({
                            "row": row_idx, 
                            "col": col_idx,
                            "value": str(cell.value),
                            "is_bold": cell.font.b or False,
                            "is_merged": cell.coordinate in [merged_range.coord for merged_range in sheet.merged_cells.ranges]
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из Excel файла.")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке Excel файла: {e}")
            return None

    def _csv_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует CSV файл в пространственный JSON."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                sample = f.read(1024)
                sniffer = csv.Sniffer()
                try:
                    dialect = sniffer.sniff(sample, delimiters=',;')
                    sep = dialect.delimiter
                except csv.Error:
                    sep = ','
            
            df = pd.read_csv(file_path, header=None, sep=sep, dtype=str)
            
            cells_data = []
            for row_idx, row in df.iterrows():
                for col_idx, value in enumerate(row):
                    if pd.notna(value) and str(value).strip():
                        cells_data.append({
                            "row": row_idx,
                            "col": col_idx,
                            "value": str(value).strip(),
                            "is_bold": False,
                            "is_merged": False
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из CSV файла.")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке CSV файла: {e}")
            return None

    def _pdf_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует PDF файл в пространственный JSON."""
        try:
            tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True, pandas_options={'header': None})
            
            if not tables:
                self.cascade_log.append("Таблицы в PDF не найдены.")
                return None
            
            combined_df = pd.concat(tables, ignore_index=True)
            
            cells_data = []
            for row_idx, row in combined_df.iterrows():
                for col_idx, value in enumerate(row):
                    if pd.notna(value) and str(value).strip():
                        cells_data.append({
                            "row": row_idx,
                            "col": col_idx,
                            "value": str(value).strip(),
                            "is_bold": False,
                            "is_merged": False
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из PDF файла ({len(tables)} таблиц).")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке PDF файла: {e}")
            return None

    def _docx_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует DOCX файл в пространственный JSON."""
        try:
            doc = docx.Document(file_path)
            
            if not doc.tables:
                self.cascade_log.append("Таблицы в DOCX не найдены.")
                return None
            
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables_data.append((len(table_data) * len(table_data[0]) if table_data else 0, table_data))
            
            if not tables_data:
                return None
            
            _, largest_table_data = max(tables_data, key=lambda item: item[0])
            
            cells_data = []
            for row_idx, row in enumerate(largest_table_data):
                for col_idx, value in enumerate(row):
                    if value and str(value).strip():
                        cells_data.append({
                            "row": row_idx,
                            "col": col_idx,
                            "value": str(value).strip(),
                            "is_bold": False,
                            "is_merged": False
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из DOCX файла.")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке DOCX файла: {e}")
            return None

    def _file_to_dataframe(self, file_path: str) -> Optional[pd.DataFrame]:
        """Универсальный метод для чтения Excel/CSV файлов в DataFrame."""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.xls':
                # Для старых .xls файлов используем xlrd
                self.cascade_log.append(f"Чтение .xls файла с помощью xlrd: {file_path}")
                # pandas.read_excel с движком xlrd
                df = pd.read_excel(file_path, engine='xlrd')
            elif ext == '.xlsx':
                # Для новых .xlsx файлов используем openpyxl (по умолчанию для pandas)
                self.cascade_log.append(f"Чтение .xlsx файла с помощью openpyxl: {file_path}")
                df = pd.read_excel(file_path, engine='openpyxl')
            elif ext == '.csv':
                self.cascade_log.append(f"Чтение .csv файла: {file_path}")
                df = pd.read_csv(file_path)
            else:
                self.cascade_log.append(f"Формат {ext} не поддерживается для чтения в DataFrame.")
                return None
            
            if df.empty:
                self.cascade_log.append(f"DataFrame пуст после чтения файла {file_path}.")
                return None
            
            self.cascade_log.append(f"Успешно прочитан файл {file_path} в DataFrame ({len(df)} строк).")
            return df
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при чтении файла {file_path} в DataFrame: {e}")
            return None

    def _get_client_items_from_llm(self, spatial_json: str) -> Optional[List[Dict]]:
        """Использует LLM для извлечения товаров из пространственного JSON."""
        self.cascade_log.append("Шаг 2: Извлечение позиций LLM из пространственного JSON.")
        prompt = PromptTemplate(
            input_variables=["spatial_data"],
            template="""
            Ты - эксперт по извлечению данных из текстовых и табличных представлений документов. 
            Твоя задача - найти в представленных данных список товарных позиций и их количество. 
            
            ПРАВИЛА:
            1. Извлекать только те позиции, которые являются ЯВНЫМИ ТОВАРАМИ. Игнорировать заголовки, подписи, общие фразы, итоги, служебные слова и т.п.
            2. Для КАЖДОЙ позиции извлечь:
               - "full_name": Полное, точное название товара, как оно написано в документе. Не сокращай, не изменяй, объединяй связанные части (например, "Фланец Ду 100 Ру 16").
               - "quantity": Числовое количество товара. Если количество не указано, используй 1. Если указано "до 50", извлекай 50.
            3. Формат ответа - ТОЛЬКО валидный JSON-массив объектов. Каждый объект должен иметь ключи "full_name" (строка) и "quantity" (целое число).
            4. Если товар имеет несколько строк (например, описание переносится на новую строку), объедини их в одно `full_name`.
            5. Обрати внимание на контекст: если столбец явно является "наименованием", а другой "количеством", используй это.
            6. Если данных недостаточно или не удалось извлечь ни одной осмысленной товарной позиции, верни пустой JSON массив: [].
            7. Избегай домысливания. Извлекай только то, что явно присутствует.
            8. Если в одной ячейке несколько товаров, раздели их на отдельные позиции.
            
            Примеры ввода (пространственный JSON):
            [{{"row":0,"col":0,"value":"Наименование","is_bold":true}},{{"row":0,"col":1,"value":"Количество","is_bold":true}},{{"row":1,"col":0,"value":"Труба стальная 102х4","is_bold":false}},{{"row":1,"col":1,"value":"10 шт","is_bold":false}}]
            Ожидаемый вывод: 
            [{{"full_name": "Труба стальная 102х4", "quantity": 10}}]

            Примеры ввода (пространственный JSON):
            [{{"row":0,"col":0,"value":"Перечень:","is_bold":true}},{{"row":1,"col":0,"value":"1. Отвод 90 градусов Ду500","is_bold":false}},{{"row":1,"col":1,"value":"2 шт","is_bold":false}},{{"row":2,"col":0,"value":"2. Фланец стальной Ду100 Ру16 ст.20","is_bold":false}},{{"row":2,"col":1,"value":"5 шт","is_bold":false}}]
            Ожидаемый вывод:
            [{{"full_name": "Отвод 90 градусов Ду500", "quantity": 2}}, {{"full_name": "Фланец стальной Ду100 Ру16 ст.20", "quantity": 5}}]

            Пример ввода (текстовое представление):
            Перечень:
            1. Отвод 90 градусов Ду500 - 2 шт
            2. Фланец стальной Ду100 Ру16 ст.20 - 5 шт
            
            Ожидаемый вывод:
            [{{"full_name": "Отвод 90 градусов Ду500", "quantity": 2}}, {{"full_name": "Фланец стальной Ду100 Ру16 ст.20", "quantity": 5}}]

            Представленные данные (JSON): {{spatial_data}}
            
            Вывод JSON:
            """
        )
        try:
            # TODO: Переделать LLMChain на новый синтаксис prompt | llm
            chain = LLMChain(llm=self.llm, prompt=prompt)
            response_text = chain.run(spatial_data=spatial_json)
            # Attempt to clean and parse JSON
            # Sometimes LLM may add extra text like ```json ... ```
            if response_text.strip().startswith('```json'):
                response_text = response_text.strip()[len('```json'):].strip()
                if response_text.endswith('```'):
                    response_text = response_text[:-len('```')].strip()
            
            items = json.loads(response_text)
            self.cascade_log.append(f"LLM извлекла: {len(items)} позиций.")
            return items
        except json.JSONDecodeError as e:
            self.cascade_log.append(f"LLM вернула невалидный JSON: {e}. Ответ LLM: {response_text[:500]}...")
            return None
        except Exception as e:
            self.cascade_log.append(f"Ошибка при запросе к LLM (извлечение позиций из пространственного JSON): {e}")
            return None

    def _get_client_items_from_text(self, text_data: str) -> Optional[List[Dict]]:
        """Использует LLM для извлечения товаров из обычного текста (для структурного анализа)."""
        self.cascade_log.append("Шаг 2: Извлечение позиций LLM из текстового представления.")
        prompt = PromptTemplate(
            input_variables=["text_data"],
            template="""
            Ты - эксперт по извлечению данных из неструктурированных текстовых представлений документов. 
            Твоя задача - найти в представленных данных список товарных позиций и их количество. 
            
            ПРАВИЛА:
            1. Извлекать только те позиции, которые являются ЯВНЫМИ ТОВАРАМИ. Игнорировать заголовки, подписи, общие фразы, итоги, служебные слова и т.п.
            2. Для КАЖДОЙ позиции извлечь:
               - "full_name": Полное, точное название товара, как оно написано в документе. Не сокращай, не изменяй, объединяй связанные части (например, "Фланец Ду 100 Ру 16").
               - "quantity": Числовое количество товара. Если количество не указано, используй 1. Если указано "до 50", извлекай 50.
            3. Формат ответа - ТОЛЬКО валидный JSON-массив объектов. Каждый объект должен иметь ключи "full_name" (строка) и "quantity" (целое число).
            4. Если товар имеет несколько строк (например, описание переносится на новую строку), объедини их в одно `full_name`.
            5. Обрати внимание на контекст: если столбец явно является "наименованием", а другой "количеством", используй это.
            6. Если данных недостаточно или не удалось извлечь ни одной осмысленной товарной позиции, верни пустой JSON массив: [].
            7. Избегай домысливания. Извлекай только то, что явно присутствует.
            8. Если в одной ячейке несколько товаров, раздели их на отдельные позиции.
            
            Примеры ввода:
            Перечень:
            1. Отвод 90 градусов Ду500 - 2 шт
            2. Фланец стальной Ду100 Ру16 ст.20 - 5 шт
            
            Ожидаемый вывод:
            [{{"full_name": "Отвод 90 градусов Ду500", "quantity": 2}}, {{"full_name": "Фланец стальной Ду100 Ру16 ст.20", "quantity": 5}}]

            Представленные данные (текст):
            {{text_data}}
            
            Вывод JSON:
            """
        )
        try:
            # TODO: Переделать LLMChain на новый синтаксис prompt | llm
            chain = LLMChain(llm=self.llm, prompt=prompt)
            response_text = chain.run(text_data=text_data)
            
            # Attempt to clean and parse JSON
            if response_text.strip().startswith('```json'):
                response_text = response_text.strip()[len('```json'):].strip()
                if response_text.endswith('```'):
                    response_text = response_text[:-len('```')].strip()

            items = json.loads(response_text)
            self.cascade_log.append(f"LLM извлекла: {len(items)} позиций.")
            return items
        except json.JSONDecodeError as e:
            self.cascade_log.append(f"LLM вернула невалидный JSON: {e}. Ответ LLM: {response_text[:500]}...")
            return None
        except Exception as e:
            self.cascade_log.append(f"Ошибка при запросе к LLM (извлечение позиций из текста): {e}")
            return None 