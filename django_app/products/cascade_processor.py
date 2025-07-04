import pandas as pd
import re
import logging
import numpy as np
from typing import List, Dict, Optional, Any, Tuple
import os
import json
import csv
import docx
import tabula # Для извлечения таблиц из PDF
import openpyxl # Для детального анализа стилей
from pydantic import BaseModel, ValidationError, Field

logger = logging.getLogger('commercial_proposal')

# --- Pydantic модели для валидации ответов LLM ---

class ColumnMap(BaseModel):
    price_col_index: Optional[int] = None
    stock_col_index: Optional[int] = None
    name_parts_col_indices: List[int] = Field(..., min_items=1)

class PriceListMap(BaseModel):
    header_row_index: int
    data_start_row_index: int
    column_map: ColumnMap

class AuditorVerdict(BaseModel):
    is_correct: bool
    reasoning: str

class ExtractedProduct(BaseModel):
    """Модель для валидации одного товара, извлеченного LLM."""
    full_name: str = Field(..., min_length=3)
    price: Optional[float] = None
    stock: str = "в наличии"

class CascadeProcessor:
    """
    Умный каскадный процессор "Хирург 3.0" для извлечения данных из прайс-листов.
    Использует пространственный анализ LLM для максимальной точности.
    """

    def __init__(self, llm: Any):
        if not llm:
            raise ValueError("LLM instance is required.")
        self.llm = llm
        self.cascade_log = []

    def process_file_cascade(self, file_path: str, file_name: str) -> Dict:
        """
        Надежная каскадная обработка файла с тремя уровнями и полной валидацией.
        
        УРОВЕНЬ 1: Пространственный анализ LLM (для сложных иерархических структур)
        УРОВЕНЬ 2: Структурный анализ LLM + извлечение (для стандартных таблиц)
        УРОВЕНЬ 3: Эвристические методы (fallback для простых случаев)
        """
        logger.info(f"--- ЗАПУСК КАСКАДНОЙ СИСТЕМЫ для файла: {file_name} ---")
        self.cascade_log = [f"Начало каскадной обработки файла: {file_name}"]
        
        try:
            # Определяем формат файла
            ext = os.path.splitext(file_path)[1].lower()
            self.cascade_log.append(f"Определен формат файла: {ext}")
            
            # Инициализируем результаты уровней по умолчанию, чтобы избежать ошибок UndefinedLocalVariable
            level1_result = {"success": False, "products": [], "error": "Уровень 1 не был запущен для данного формата или не дал результата."}
            level2_result = {"success": False, "products": [], "error": "Уровень 2 не был запущен или не дал результата."}
            level3_result = {"success": False, "products": [], "error": "Уровень 3 не был запущен или не дал результата."}

            # === УРОВЕНЬ 1: Пространственный анализ LLM ===
            if ext in ['.xlsx', '.xls']:
                self.cascade_log.append("🔥 УРОВЕНЬ 1: Пространственный анализ LLM")
                level1_result = self._process_level1_spatial(file_path, file_name)
                
                if level1_result["success"] and len(level1_result["products"]) >= 5:
                    self.cascade_log.append(f"✅ УРОВЕНЬ 1 УСПЕШЕН: {len(level1_result['products'])} товаров")
                    level1_result["final_method"] = "Level 1: LLM Spatial Analysis"
                    level1_result["cascade_log"] = self.cascade_log
                    return level1_result
            else:
                    self.cascade_log.append(f"⚠️ УРОВЕНЬ 1: {len(level1_result.get('products', []))} товаров (недостаточно)")
                    # Если формат не поддерживается Уровнем 1 или не дал достаточного результата,
                    # level1_result уже содержит соответствующую ошибку благодаря начальной инициализации.
                    # Нет необходимости в дополнительном else-блоке здесь.
            
            # Если дошли сюда, значит УРОВЕНЬ 1 либо не был запущен (неподдерживаемый формат),
            # либо не дал достаточного количества товаров. level1_result уже имеет соответствующее
            # дефолтное или частично обновленное значение.

            # === УРОВЕНЬ 2: Структурный анализ LLM ===
            self.cascade_log.append("🔥 УРОВЕНЬ 2: Структурный анализ LLM")
            level2_result = self._process_level2_structural(file_path, file_name)
            
            if level2_result["success"] and len(level2_result["products"]) >= 3:
                self.cascade_log.append(f"✅ УРОВЕНЬ 2 УСПЕШЕН: {len(level2_result['products'])} товаров")
                level2_result["final_method"] = "Level 2: LLM Structural Analysis"
                level2_result["cascade_log"] = self.cascade_log
                return level2_result
            else:
                self.cascade_log.append(f"⚠️ УРОВЕНЬ 2: {len(level2_result.get('products', []))} товаров")
            
            # === УРОВЕНЬ 3: ЭВРИСТИЧЕСКИЕ МЕТОДЫ ===
            self.cascade_log.append("🔥 УРОВЕНЬ 3: Эвристические методы (fallback)")
            level3_result = self._process_level3_heuristics(file_path, file_name)
            
            if level3_result["success"] and len(level3_result["products"]) > 0:
                self.cascade_log.append(f"✅ УРОВЕНЬ 3 УСПЕШЕН: {len(level3_result['products'])} товаров")
                level3_result["final_method"] = "Level 3: Heuristic Methods"
                level3_result["cascade_log"] = self.cascade_log
                return level3_result
            else:
                self.cascade_log.append(f"❌ УРОВЕНЬ 3: {len(level3_result.get('products', []))} товаров")
            
            # === ВЫБОР ЛУЧШЕГО РЕЗУЛЬТАТА ===
            results = []
            if level1_result.get("success"): results.append(("Level 1", level1_result))
            if level2_result.get("success"): results.append(("Level 2", level2_result))
            if level3_result.get("success"): results.append(("Level 3", level3_result))
            
            if results:
                # Выбираем результат с наибольшим количеством товаров
                best_name, best_result = max(results, key=lambda x: len(x[1].get("products", [])))
                self.cascade_log.append(f"✅ ВЫБРАН ЛУЧШИЙ РЕЗУЛЬТАТ: {best_name} с {len(best_result['products'])} товарами")
                best_result["final_method"] = f"Best of All: {best_name}"
                best_result["cascade_log"] = self.cascade_log
                return best_result
            
            # Если все методы провалились
            error_msg = "Все уровни каскадной обработки не смогли извлечь данные из файла."
            return self._handle_error(error_msg)
            
        except Exception as e:
            return self._handle_error(f"Критическая ошибка в каскадной системе: {e}", exc_info=True)

    def _process_level1_spatial(self, file_path: str, file_name: str) -> Dict:
        """УРОВЕНЬ 1: Пространственный анализ LLM для сложных иерархических структур."""
        try:
            # Шаг 1: Преобразование файла в "пространственный" JSON
            spatial_json = self._file_to_spatial_json(file_path)
            if not spatial_json:
                return {"success": False, "products": [], "error": "Не удалось создать пространственное представление"}

            # Шаг 2: Получение готового списка товаров от LLM
            products_json = self._get_products_from_llm(spatial_json)
            if not products_json:
                return {"success": False, "products": [], "error": "LLM не смогла извлечь данные из пространственного JSON"}

            # Шаг 3: Валидация и очистка данных
            validated_products = self._validate_and_clean_products(products_json, "Level 1")
            
            if not validated_products:
                return {"success": False, "products": [], "error": "Ни один товар не прошел валидацию"}

            return {"success": True, "products": validated_products}
            
        except Exception as e:
            return {"success": False, "products": [], "error": f"Ошибка уровня 1: {e}"}

    def _process_level2_structural(self, file_path: str, file_name: str) -> Dict:
        """УРОВЕНЬ 2: Структурный анализ LLM для стандартных таблиц."""
        try:
            # Шаг 1: Преобразование файла в DataFrame
            df = self._file_to_dataframe(file_path)
            if df is None:
                return {"success": False, "products": [], "error": "Не удалось прочитать файл"}

            # Шаг 2: Получение карты структуры от LLM
            structure_map, sample_text = self._get_structure_map_from_llm(df)
            if not structure_map:
                return {"success": False, "products": [], "error": "LLM не смогла определить структуру файла"}

            # Шаг 3: Валидация карты структуры
            try:
                validated_map = PriceListMap(**structure_map)
            except ValidationError as e:
                return {"success": False, "products": [], "error": f"Карта структуры не прошла валидацию: {e}"}

            # Шаг 4: Проверка карты аудитором (опционально)
            if self.llm:
                verdict = self._get_auditor_verdict(sample_text, structure_map)
                if verdict and not verdict.is_correct:
                    self.cascade_log.append(f"Аудитор отклонил карту: {verdict.reasoning}")
                    # Продолжаем, но с предупреждением

            # Шаг 5: Извлечение товаров по карте
            raw_products = self._extract_products_with_map(df, validated_map)
            if not raw_products:
                return {"success": False, "products": [], "error": "Не удалось извлечь товары по карте структуры"}

            # Шаг 6: Валидация и очистка данных
            validated_products = self._validate_and_clean_products(raw_products, "Level 2")
            
            if not validated_products:
                return {"success": False, "products": [], "error": "Ни один товар не прошел валидацию"}

            return {"success": True, "products": validated_products}
            
        except Exception as e:
            return {"success": False, "products": [], "error": f"Ошибка уровня 2: {e}"}

    def _process_level3_heuristics(self, file_path: str, file_name: str) -> Dict:
        """УРОВЕНЬ 3: Эвристические методы (fallback)."""
        try:
            result = self._process_with_heuristics(file_path)
            if result["success"]:
                # Валидация и очистка данных
                validated_products = self._validate_and_clean_products(result["products"], "Level 3")
                return {"success": True, "products": validated_products}
            else:
                return result
                
        except Exception as e:
            return {"success": False, "products": [], "error": f"Ошибка уровня 3: {e}"}

    def _validate_and_clean_products(self, products: List[Dict], level_name: str) -> List[Dict]:
        """
        Универсальная валидация и очистка товаров для всех уровней.
        
        Выполняет:
        1. Pydantic валидацию
        2. Проверку качества данных
        3. Очистку и нормализацию
        4. Фильтрацию дубликатов
        """
        validated_products = []
        skipped_count = 0
        
        for i, item in enumerate(products):
            try:
                # Нормализация структуры данных
                normalized_item = self._normalize_product_structure(item)
                
                # Pydantic валидация
                validated_product = ExtractedProduct(**normalized_item)
                
                # Проверка качества данных
                if self._is_quality_product(validated_product):
                    validated_products.append(validated_product.dict())
                else:
                    skipped_count += 1
                    self.cascade_log.append(f"{level_name}: Пропущен товар низкого качества: {normalized_item.get('full_name', 'Без названия')[:50]}")
                    
            except ValidationError as e:
                skipped_count += 1
                self.cascade_log.append(f"{level_name}: Ошибка валидации товара {i+1}: {e}")
                continue
            except Exception as e:
                skipped_count += 1
                self.cascade_log.append(f"{level_name}: Неожиданная ошибка при валидации товара {i+1}: {e}")
                continue
        
        # Удаление дубликатов
        unique_products = self._remove_duplicates(validated_products)
        
        removed_duplicates = len(validated_products) - len(unique_products)
        if removed_duplicates > 0:
            self.cascade_log.append(f"{level_name}: Удалено {removed_duplicates} дубликатов")
        
        self.cascade_log.append(f"{level_name}: Валидация завершена. Принято: {len(unique_products)}, Отклонено: {skipped_count}")
        
        return unique_products

    def _normalize_product_structure(self, item: Dict) -> Dict:
        """Нормализует структуру данных товара для унификации между уровнями."""
        normalized = {
            "full_name": "",
            "price": None,
            "stock": "в наличии"
        }
        
        # Извлекаем название
        if "full_name" in item:
            normalized["full_name"] = str(item["full_name"]).strip()
        elif "name" in item:
            normalized["full_name"] = str(item["name"]).strip()
        
        # Извлекаем цену
        if "price" in item and item["price"] is not None:
            try:
                if isinstance(item["price"], (int, float)):
                    normalized["price"] = float(item["price"])
                else:
                    # Пытаемся очистить и преобразовать строку
                    cleaned_price = self._clean_price(item["price"])
                    normalized["price"] = cleaned_price
            except:
                normalized["price"] = None
        
        # Извлекаем остаток
        if "stock" in item and item["stock"] is not None:
            normalized["stock"] = self._clean_stock(item["stock"])
        
        return normalized

    def _is_quality_product(self, product: ExtractedProduct) -> bool:
        """
        Проверяет качество товара по множественным критериям.
        
        Критерии качества:
        1. Название не менее 3 символов
        2. Название не является служебным словом
        3. Цена разумная (если указана)
        4. Название содержит значимую информацию
        """
        name = product.full_name.lower().strip()
        
        # Проверка длины
        if len(name) < 3:
            return False
        
        # Проверка на служебные слова
        service_words = [
            'nan', 'none', 'null', 'undefined', 'наименование', 'товар', 'продукт',
            'название', 'описание', 'итого', 'всего', 'сумма', 'total', 'sum',
            'заголовок', 'header', 'title', 'примечание', 'note', 'комментарий'
        ]
        
        if name in service_words:
            return False
        
        # Проверка на слишком короткие или бессмысленные названия
        if len(name.split()) < 2 and not any(char.isdigit() for char in name):
            # Исключение: если название содержит цифры, оно может быть коротким (например, "Кран DN50")
            return False
        
        # Проверка цены (если указана)
        if product.price is not None:
            if product.price < 0 or product.price > 10000000:  # Разумные пределы
                return False
        
        # Проверка на минимальное содержание информации
        if len(name.replace(' ', '')) < 3:
            return False
        
        return True

    def _remove_duplicates(self, products: List[Dict]) -> List[Dict]:
        """Удаляет ИСТИННЫЕ дубликаты товаров (одинаковые название + цена + остаток)."""
        seen_combinations = set()
        unique_products = []
        
        for product in products:
            # Создаем ключ из названия, цены и остатка
            name_key = product["full_name"].lower().strip()
            price_key = str(product.get("price", ""))
            stock_key = str(product.get("stock", ""))
            
            # Комбинированный ключ для проверки истинных дубликатов
            combination_key = f"{name_key}|{price_key}|{stock_key}"
            
            if combination_key not in seen_combinations:
                seen_combinations.add(combination_key)
                unique_products.append(product)
        
        return unique_products

    def _handle_error(self, message: str, exc_info=False) -> Dict:
        """Централизованный обработчик ошибок."""
        logger.error(message, exc_info=exc_info)
        self.cascade_log.append(f"ОШИБКА: {message}")
        return {"success": False, "products": [], "error": message, "cascade_log": self.cascade_log}

    def _file_to_spatial_json(self, file_path: str) -> Optional[str]:
        """
        Преобразует файл любого поддерживаемого формата в детализированный JSON 
        с координатами и стилями для пространственного анализа.
        """
        self.cascade_log.append("Шаг 1: Создание пространственного JSON представления.")
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext in ['.xlsx', '.xls']:
                # Excel файлы - полный пространственный анализ с координатами и стилями
                return self._excel_to_spatial_json(file_path)
            
            elif ext == '.csv':
                # CSV файлы - упрощенный пространственный анализ
                return self._csv_to_spatial_json(file_path)
            
            elif ext == '.pdf':
                # PDF файлы - извлечение таблиц и преобразование в пространственный JSON
                return self._pdf_to_spatial_json(file_path)
            
            elif ext == '.docx':
                # DOCX файлы - извлечение таблиц и преобразование в пространственный JSON
                return self._docx_to_spatial_json(file_path)
            
            else:
                self.cascade_log.append(f"Формат {ext} не поддерживается для пространственного анализа.")
                return None
                
        except Exception as e:
            self.cascade_log.append(f"Ошибка при создании пространственного JSON: {e}")
            return None

    def _excel_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует Excel файл в пространственный JSON."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            sheet = workbook.active
            
            cells_data = []
            # Ограничиваем количество строк для анализа, чтобы не превысить лимиты токенов
            max_rows_to_process = 500  # Увеличено для извлечения всех товаров из файлов типа УралОтвод
            for row_idx, row in enumerate(sheet.iter_rows(max_row=max_rows_to_process)):
                for col_idx, cell in enumerate(row):
                    if cell.value is not None:
                        cells_data.append({
                            "row": row_idx, 
                            "col": col_idx,
                            "value": str(cell.value),
                            "is_bold": cell.font.b or False,
                            "is_merged": cell.coordinate in [merged_range.coord for merged_range in sheet.merged_cells.ranges]
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из Excel файла.")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке Excel файла: {e}")
            return None

    def _csv_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует CSV файл в пространственный JSON."""
        try:
            # Пробуем определить разделитель
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                sample = f.read(1024)
                sniffer = csv.Sniffer()
                try:
                    dialect = sniffer.sniff(sample, delimiters=',;')
                    sep = dialect.delimiter
                except csv.Error:
                    sep = ','
            
            # Читаем CSV
            df = pd.read_csv(file_path, header=None, sep=sep, dtype=str)
            
            cells_data = []
            for row_idx, row in df.iterrows():
                for col_idx, value in enumerate(row):
                    if pd.notna(value) and str(value).strip():
                        cells_data.append({
                            "row": row_idx,
                            "col": col_idx,
                            "value": str(value).strip(),
                            "is_bold": False,  # CSV не поддерживает форматирование
                            "is_merged": False
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из CSV файла.")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке CSV файла: {e}")
            return None

    def _pdf_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует PDF файл в пространственный JSON."""
        try:
            # Извлекаем таблицы из PDF
            tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True, pandas_options={'header': None})
            
            if not tables:
                self.cascade_log.append("Таблицы в PDF не найдены.")
                return None
            
            # Объединяем все таблицы
            combined_df = pd.concat(tables, ignore_index=True)
            
            cells_data = []
            for row_idx, row in combined_df.iterrows():
                for col_idx, value in enumerate(row):
                    if pd.notna(value) and str(value).strip():
                        cells_data.append({
                            "row": row_idx,
                            "col": col_idx,
                            "value": str(value).strip(),
                            "is_bold": False,  # PDF извлечение не сохраняет форматирование
                            "is_merged": False
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из PDF файла ({len(tables)} таблиц).")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке PDF файла: {e}")
            return None

    def _docx_to_spatial_json(self, file_path: str) -> Optional[str]:
        """Преобразует DOCX файл в пространственный JSON."""
        try:
            doc = docx.Document(file_path)
            
            if not doc.tables:
                self.cascade_log.append("Таблицы в DOCX не найдены.")
                return None
            
            # Находим самую большую таблицу
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables_data.append((len(table_data) * len(table_data[0]) if table_data else 0, table_data))
            
            if not tables_data:
                return None
            
            # Выбираем самую большую таблицу
            _, largest_table_data = max(tables_data, key=lambda item: item[0])
            
            cells_data = []
            for row_idx, row in enumerate(largest_table_data):
                for col_idx, value in enumerate(row):
                    if value and str(value).strip():
                        cells_data.append({
                            "row": row_idx,
                            "col": col_idx,
                            "value": str(value).strip(),
                            "is_bold": False,  # Упрощенное извлечение без форматирования
                            "is_merged": False
                        })
            
            self.cascade_log.append(f"Проанализировано {len(cells_data)} ячеек из DOCX файла.")
            return json.dumps(cells_data, ensure_ascii=False)
            
        except Exception as e:
            self.cascade_log.append(f"Ошибка при обработке DOCX файла: {e}")
            return None

    def _get_products_from_llm(self, spatial_json: str) -> Optional[List[Dict]]:
        """Отправляет пространственный JSON в LLM и получает готовый список товаров."""
        self.cascade_log.append("Шаг 2: Запрос на извлечение товаров у LLM.")
        prompt = self._get_spatial_analyst_prompt(spatial_json)
        
        try:
            response = self.llm.invoke(prompt)
            response_text = response.content if hasattr(response, 'content') else str(response)
            self.cascade_log.append("Ответ от LLM получен.")
            
            # Извлекаем JSON из ответа, который может быть обернут в markdown
            match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if match:
                json_str = match.group(0)
                return json.loads(json_str)
            self.cascade_log.append("JSON-массив не найден в ответе LLM.")
            return None
        except Exception as e:
            self.cascade_log.append(f"Ошибка при получении или парсинге ответа от LLM: {e}")
            return None

    def _get_spatial_analyst_prompt(self, spatial_json: str) -> str:
        return f"""
Ты — эксперт по анализу сложных прайс-листов. Перед тобой JSON-представление файла с данными, где для каждой ячейки указаны ее координаты (`row`, `col`), значение (`value`) и стиль (`is_bold`, `is_merged`).

Твоя задача — проанализировать эту структуру и вернуть **готовый, собранный список товаров.**

**КРИТИЧЕСКИ ВАЖНЫЕ ИНСТРУКЦИИ:**

1. **Найди иерархию и структуру:**
   - Определи заголовки групп (например, "Задвижки чугунные", "Фланцы стальные") - они часто выделены жирным (`is_bold: true`) или находятся в отдельных строках
   - Найди строку с заголовками колонок (Наименование, Цена, Остаток, Ду, Ру и т.д.)
   - Определи где начинаются данные товаров

2. **Собери полные названия товаров:**
   - Для каждого товара собери ПОЛНОЕ, осмысленное наименование
   - Название может состоять из: `Заголовок группы` + `Базовое имя` + `Характеристики` (Ду, Ру, ГОСТ, материал и т.д.)
   - Пример: "Задвижки чугунные" + "30ч6бр" + "Ду 50" + "Ру 10" = "Задвижки чугунные 30ч6бр Ду 50 Ру 10"

3. **Извлеки цены и остатки:**
   - Найди колонки с ценами (числовые значения, обычно 3-7 цифр)
   - Найди колонки с остатками (могут содержать числа или текст типа "в наличии", "под заказ")
   - Если цена не найдена, используй `null`
   - Если остаток не найден, используй "в наличии"

4. **Фильтрация:**
   - НЕ включай строки заголовков, пустые строки, итоговые строки
   - НЕ включай служебную информацию (номера страниц, примечания и т.д.)
   - Включай только реальные товары с осмысленными названиями

5. **Обработка цен:**
   - Цены могут содержать пробелы, запятые как разделители тысяч
   - Преобразуй в числовой формат (например: "1 234,56" → 1234.56)
   - Если цену невозможно распознать, используй `null`

**Вот JSON-представление файла:**
---
{spatial_json}
---

**ФОРМАТ ОТВЕТА - ТОЛЬКО JSON-массив без комментариев:**
```json
[
  {{
    "full_name": "Задвижка чугунная 30ч6бр Ду 50 Ру 10",
    "price": 3600.00,
    "stock": "в наличии"
  }},
  {{
    "full_name": "Фланец стальной плоский Ду 100 Ру 16 ГОСТ 12820",
    "price": 1250.50,
    "stock": "под заказ"
  }}
]
```

ВАЖНО: Анализируй координаты ячеек для понимания структуры. Товары обычно находятся в одной строке, но название может быть составлено из нескольких колонок."""

    def _file_to_dataframe(self, file_path: str) -> Optional[pd.DataFrame]:
        """
        Шаг -1: "Универсальный Пре-процессор".
        Преобразует файл любого поддерживаемого формата в pandas DataFrame.
        """
        self.cascade_log.append("Шаг -1: Универсальное чтение файла.")
        ext = os.path.splitext(file_path)[1].lower()
        self.cascade_log.append(f"Определен формат: {ext}")
        
        try:
            if ext in ['.xlsx', '.xls', '.xlsm', '.xlsb', '.odf', '.ods', '.odt']:
                return pd.read_excel(file_path, header=None, sheet_name=0)
            
            elif ext == '.csv':
                # Пробуем определить разделитель
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    sample = f.read(1024)
                    sniffer = csv.Sniffer()
                    try:
                        dialect = sniffer.sniff(sample, delimiters=',;')
                        sep = dialect.delimiter
                    except csv.Error:
                        sep = ',' # По умолчанию
                return pd.read_csv(file_path, header=None, sep=sep)

            elif ext == '.pdf':
                self.cascade_log.append("Попытка извлечения таблиц из PDF с помощью Tabula...")
                # Tabula извлекает список DataFrame'ов, по одному для каждой таблицы
                tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True, pandas_options={'header': None})
                if tables:
                    self.cascade_log.append(f"Найдено {len(tables)} таблиц в PDF. Объединяем.")
                    return pd.concat(tables, ignore_index=True)
                else:
                    self.cascade_log.append("Таблицы в PDF не найдены с помощью Tabula.")
                    return None

            elif ext == '.docx':
                self.cascade_log.append("Попытка извлечения таблиц из DOCX...")
                doc = docx.Document(file_path)
                if doc.tables:
                    # Берем самую большую таблицу из документа
                    tables_data = []
                    for table in doc.tables:
                        table_data = []
                        for row in table.rows:
                            table_data.append([cell.text for cell in row.cells])
                        tables_data.append((len(table_data) * len(table_data[0]) if table_data else 0, table_data))
                    
                    if tables_data:
                        # Выбираем самую большую таблицу по количеству ячеек
                        _, largest_table_data = max(tables_data, key=lambda item: item[0])
                        return pd.DataFrame(largest_table_data)
                else:
                    self.cascade_log.append("Таблицы в DOCX не найдены.")
                    return None
            
            else:
                self.cascade_log.append(f"Неподдерживаемый формат файла для извлечения прайс-листа: {ext}")
                return None
        except Exception as e:
            logger.error(f"Ошибка при преобразовании файла '{file_path}' в DataFrame: {e}", exc_info=True)
            self.cascade_log.append(f"Ошибка при чтении файла: {e}")
            return None

    def _get_structure_map_from_llm(self, df: pd.DataFrame) -> Tuple[Optional[Dict], str]:
        """Этап "Разведчик". Анализирует структуру DataFrame с помощью LLM-Аналитика."""
        self.cascade_log.append("Этап 1: Запрос карты структуры у LLM-Аналитика.")
        try:
            # Шаг 0: Адаптивная Оценка
            if len(df) <= 200:
                self.cascade_log.append("Файл 'маленький' (<= 200 строк). Отправляется на анализ целиком.")
                sample_df = df
            else:
                self.cascade_log.append("Файл 'большой' (> 200 строк). Создается репрезентативная выборка.")
                # Убираем полностью пустые строки из выборки, чтобы не засорять контекст
                df_non_empty = df.dropna(how='all')
                sample_df = pd.concat([
                    df_non_empty.head(30), 
                    df_non_empty.sample(n=min(30, len(df_non_empty)), random_state=1), 
                    df_non_empty.tail(30)
                ]).drop_duplicates()

            sample_text = sample_df.to_csv(index=False, header=False, lineterminator='\n')
            prompt = self._get_llm_analyst_prompt(sample_text)
            
            self.cascade_log.append("Отправка запроса к LLM-Аналитику...")
            response = self.llm.invoke(prompt)
            response_text = response.content if hasattr(response, 'content') else str(response)
            self.cascade_log.append("Ответ от LLM-Аналитика получен.")

            json_map = self._parse_llm_response(response_text)
            
            # Валидация карты
            if not json_map or 'header_row_index' not in json_map or 'column_map' not in json_map:
                self.cascade_log.append(f"Ошибка: LLM вернула неполную карту структуры: {json_map}")
                return None, ""

            self.cascade_log.append(f"Карта структуры успешно разобрана: {json_map}")
            return json_map, sample_text

        except Exception as e:
            logger.error(f"Ошибка на этапе анализа LLM: {e}", exc_info=True)
            self.cascade_log.append(f"Ошибка на этапе анализа LLM-Аналитика: {e}")
            return None, ""

    def _get_llm_analyst_prompt(self, sample_text: str) -> str:
        return f"""
Ты — эксперт по анализу прайс-листов со сложной структурой. Перед тобой фрагмент прайс-листа в формате CSV. Твоя задача — максимально точно определить его структуру и вернуть JSON-объект.

**Инструкции:**
1. Номера колонок и строк начинаются с 0.
2. `header_row_index` - это номер строки, где находятся основные заголовки.
3. `data_start_row_index` - это номер строки, с которой начинаются РЕАЛЬНЫЕ товары.
4. `price_col_index` - это ИНДЕКС колонки с ценой. Она может называться "Цена", "Стоимость", "Цена с НДС", "руб.". Ищи колонку с числовыми значениями, похожими на деньги.
5. `stock_col_index` - это ИНДЕКС колонки с остатком. Она может называться "Остаток", "Наличие", "Кол-во", "Склад", "Qty".
6. `name_parts_col_indices` - это список ИНДЕКСОВ колонок, которые нужно объединить для формирования полного названия товара.
7. **ВАЖНО:** Некоторые прайсы имеют иерархическую структуру. Если ты видишь, что одна строка (например, "Задвижки чугунные") является заголовком для нескольких последующих, ты должен это учесть.

**Пример анализа сложной структуры:**
Исходные данные:
```
,,,,
,Задвижки чугунные,,,
,Наименование,Ду,Цена,Остаток
,30ч6бр,50,3500,10
,30ч6бр,80,4500,12
```
Ожидаемый JSON:
{{
  "header_row_index": 2,
  "data_start_row_index": 3,
  "column_map": {{
    "price_col_index": 3,
    "stock_col_index": 4,
    "name_parts_col_indices": [1, 2]
  }}
}}
*Объяснение: Заголовок группы "Задвижки чугунные" находится выше основной таблицы, поэтому для сборки полного имени нужно будет использовать логику иерархии в последующем шаге.*

**Фрагмент для анализа:**
---
{sample_text}
---

Верни ТОЛЬКО JSON-объект. Не добавляй никаких комментариев вне JSON.
"""

    def _parse_llm_response(self, response_text: str) -> Optional[Dict]:
        """Извлекает и парсит JSON из текстового ответа LLM."""
        try:
            match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if match:
                json_str = match.group(0)
                return json.loads(json_str)
            else:
                self.cascade_log.append("JSON не найден в ответе LLM.")
                return None
        except json.JSONDecodeError as e:
            self.cascade_log.append(f"Ошибка декодирования JSON из ответа LLM: {e}")
            return None

    def _get_auditor_verdict(self, sample_text: str, analysis_map: Dict) -> Optional[AuditorVerdict]:
        """Шаг 'Консилиум'. Аудитор (второй LLM) проверяет карту Аналитика."""
        self.cascade_log.append("Шаг 3: Запрос вердикта у LLM-Аудитора.")
        
        prompt = f"""
Ты — старший аудитор данных. Твоя задача — проверить работу младшего аналитика. Ниже представлены исходные данные и JSON-карта, которую он сгенерировал.

**Исходные данные (фрагмент CSV):**
---
{sample_text}
---
**JSON-карта от аналитика:**
---
{json.dumps(analysis_map, indent=2, ensure_ascii=False)}
---
**Твоя задача:**
Проверь JSON-карту на ЛОГИЧЕСКУЮ корректность.
1. Убедись, что `price_col_index` действительно указывает на колонку с ценами, а не с чем-то другим (артикулы, ГОСТ, количество). Допускается, если в колонке есть пропуски или текстовые значения, но основное содержимое должно быть похоже на цены.
2. Убедись, что `name_parts_col_indices` включает в себя все колонки, необходимые для полного и осмысленного названия.

**Верни ТОЛЬКО JSON-объект с твоим вердиктом:**
{{
  "is_correct": <true/false>,
  "reasoning": "<краткое текстовое обоснование, особенно если is_correct: false. Если все верно, напиши 'Карта выглядит логичной и корректной.'>"
}}
"""
        try:
            response = self.llm.invoke(prompt)
            response_text = response.content if hasattr(response, 'content') else str(response)
            verdict_json = self._parse_llm_response(response_text)
            if verdict_json:
                return AuditorVerdict(**verdict_json)
            return None
        except (ValidationError, Exception) as e:
            self.cascade_log.append(f"Ошибка при получении или валидации вердикта аудитора: {e}")
            return None

    def _extract_products_with_map(self, df: pd.DataFrame, structure_map: PriceListMap) -> List[Dict]:
        """Этап "Хирург". Извлекает товары из DataFrame, используя ВАЛИДНУЮ карту структуры."""
        self.cascade_log.append("Шаг 5: Начало извлечения товаров по валидной карте.")
        
        header_row = structure_map.header_row_index
        data_start_row = structure_map.data_start_row_index
        
        col_map = structure_map.column_map
        name_parts_cols_indices = col_map.name_parts_col_indices
        price_col_index = col_map.price_col_index
        stock_col_index = col_map.stock_col_index
        
        # Данные начинаются со строки data_start_row
        data_df = df.iloc[data_start_row:].reset_index(drop=True)
        
        products = []
        current_group_name = ""
        current_subgroup_name = ""
        
        # Ищем заголовки групп в строках выше data_start_row
        group_headers = self._find_group_headers(df, header_row, data_start_row)
        self.cascade_log.append(f"Найдено {len(group_headers)} заголовков групп: {list(group_headers.values())}")

        for index, row in data_df.iterrows():
            # Проверяем, является ли строка заголовком подгруппы
            if self._is_subgroup_header(row, name_parts_cols_indices, price_col_index):
                # Это заголовок подгруппы
                subgroup_parts = []
                for col_idx in name_parts_cols_indices:
                    if col_idx < len(row) and pd.notna(row.iloc[col_idx]):
                        part = str(row.iloc[col_idx]).strip()
                        if part and part.lower() not in ['nan', 'none', '']:
                            subgroup_parts.append(part)
                
                if subgroup_parts:
                    current_subgroup_name = " ".join(subgroup_parts)
                    self.cascade_log.append(f"Обнаружен заголовок подгруппы: '{current_subgroup_name}'")
                continue
            
            # Собираем полное наименование из частей по индексам
            name_parts = []
            for col_idx in name_parts_cols_indices:
                if col_idx < len(row) and pd.notna(row.iloc[col_idx]):
                    part = str(row.iloc[col_idx]).strip()
                    if part and part.lower() not in ['nan', 'none', '']:
                        name_parts.append(part)
            
            if not name_parts:
                continue
            
            # Определяем группу по позиции строки
            actual_row_index = data_start_row + index
            group_name = self._get_group_for_row(actual_row_index, group_headers)
            
            # Собираем полное название с учетом сложной структуры УАЗ
            full_name_parts = []
            
            # Добавляем название группы
            if group_name:
                full_name_parts.append(group_name)
            
            # Добавляем название подгруппы (если это не просто число)
            if current_subgroup_name and current_subgroup_name not in group_name:
                if not current_subgroup_name.isdigit():  # Не добавляем просто числа как подгруппы
                    full_name_parts.append(current_subgroup_name)
            
            # Добавляем части названия из строки
            model_name = name_parts[0] if name_parts else ""
            diameter = name_parts[1] if len(name_parts) > 1 else ""
            
            # Добавляем модель (если это не просто число и не содержит уже "Ру")
            if model_name and not model_name.isdigit():
                # Проверяем, не содержит ли модель уже информацию о Ру
                if 'ру' not in model_name.lower():
                    full_name_parts.append(model_name)
                else:
                    # Если содержит Ру, добавляем как есть
                    full_name_parts.append(model_name)
            elif model_name and model_name.isdigit():
                # Если model_name это число, это скорее всего диаметр
                diameter = model_name
                model_name = ""
            
            # Добавляем диаметр с префиксом "Ду"
            if diameter and diameter.isdigit():
                full_name_parts.append(f"Ду {diameter}")
            elif diameter and not diameter.isdigit():
                # Если диаметр не число, но есть значение, добавляем как есть
                full_name_parts.append(diameter)
            
            # Ищем информацию о давлении (Ру) в строке заголовков
            ru_info = self._find_ru_info(df, actual_row_index, col_map.price_col_index)
            if ru_info and not any('ру' in part.lower() for part in full_name_parts):
                full_name_parts.append(ru_info)
            
            full_name = " ".join(full_name_parts).strip()
            
            # Извлекаем цену по индексу
            price = None
            if price_col_index is not None and price_col_index < len(row):
                price = self._clean_price(row.iloc[price_col_index])

            if not full_name or len(full_name) < 3:
                continue

            # Извлекаем остаток по индексу
            stock = 'в наличии'
            if stock_col_index is not None and stock_col_index < len(row):
                stock = self._clean_stock(row.iloc[stock_col_index])
            else:
                # Если колонка остатка не найдена, ищем ее "по смыслу" в строке
                stock_val_from_row = next((str(v) for v in row if isinstance(v, str) and any(w in v.lower() for w in ['наличи', 'заказ'])), 'в наличии')
                stock = self._clean_stock(stock_val_from_row)

            products.append({"name": full_name, "price": price, "stock": stock})
            
        self.cascade_log.append(f"Извлечено {len(products)} товаров с полными названиями")
        return products

    def _find_group_headers(self, df: pd.DataFrame, header_row: int, data_start_row: int) -> Dict[int, str]:
        """Находит заголовки групп во всем файле."""
        group_headers = {}
        
        # Ищем заголовки групп во всем файле
        for row_idx in range(len(df)):
            if row_idx == header_row:  # Пропускаем строку с заголовками колонок
                continue
                
            row = df.iloc[row_idx]
            
            # Ищем ячейки с текстом, которые могут быть заголовками групп
            for col_idx, value in enumerate(row):
                if pd.notna(value) and isinstance(value, str):
                    value = str(value).strip()
                    # Расширенный поиск заголовков групп
                    if (len(value) > 8 and  # Увеличили минимальную длину
                        any(keyword in value.lower() for keyword in [
                            'задвижк', 'фланец', 'отвод', 'тройник', 'переход', 'клапан', 
                            'кран', 'затвор', 'вентил', 'фильтр', 'муфта', 'чугун', 'сталь',
                            'арматура', 'трубопровод', 'соединение', 'крепеж', 'болт', 'гайка',
                            'шайба', 'прокладка', 'уплотнение', 'редуктор', 'насос', 'компенсатор',
                            'опора', 'подвеска', 'изоляция', 'теплоизоляция', 'цепь', 'канат',
                            'строп', 'такелаж', 'грузоподъем'
                        ]) and
                        value.lower() not in ['наименование', 'цена', 'остаток', 'артикул', 'гост', 'ту']):
                        group_headers[row_idx] = value
                        self.cascade_log.append(f"Найден заголовок группы в строке {row_idx}: '{value}'")
                        break
        
        return group_headers

    def _get_group_for_row(self, row_index: int, group_headers: Dict[int, str]) -> str:
        """Определяет, к какой группе относится строка с товаром."""
        # Находим ближайший заголовок группы выше текущей строки
        applicable_headers = [(header_row, name) for header_row, name in group_headers.items() if header_row < row_index]
        
        if applicable_headers:
            # Берем самый близкий заголовок
            _, group_name = max(applicable_headers, key=lambda x: x[0])
            return group_name
        
        return ""

    def _is_subgroup_header(self, row: pd.Series, name_cols: List[int], price_col: Optional[int]) -> bool:
        """Проверяет, является ли строка заголовком подгруппы."""
        # Заголовок подгруппы обычно имеет название, но не имеет цены
        has_name = False
        has_price = False
        
        # Проверяем наличие названия
        for col_idx in name_cols:
            if col_idx < len(row) and pd.notna(row.iloc[col_idx]):
                value = str(row.iloc[col_idx]).strip()
                if value and value.lower() not in ['nan', 'none', '']:
                    has_name = True
                    break
        
        # Проверяем отсутствие цены
        if price_col is not None and price_col < len(row):
            price_val = row.iloc[price_col]
            if pd.notna(price_val):
                price_str = str(price_val).strip()
                # Если в колонке цены есть число, это не заголовок
                if re.search(r'\d+', price_str):
                    has_price = True
        
        # Заголовок подгруппы: есть название, но нет цены
        return has_name and not has_price

    def _find_ru_info(self, df: pd.DataFrame, row_index: int, price_col_index: Optional[int]) -> str:
        """Находит информацию о давлении (Ру) для товара по колонке цены."""
        if price_col_index is None:
            return ""
        
        # Ищем в строках выше текущей строки информацию о Ру
        for check_row in range(max(0, row_index - 5), row_index):
            if check_row < len(df):
                row = df.iloc[check_row]
                if price_col_index < len(row) and pd.notna(row.iloc[price_col_index]):
                    value = str(row.iloc[price_col_index]).strip()
                    if 'ру' in value.lower() and any(char.isdigit() for char in value):
                        return value
        
        return ""

    def _clean_price(self, price_val: Any) -> Optional[float]:
        if pd.isna(price_val): return None
        try:
            price_str = re.sub(r'[^\d,.]', '', str(price_val)).replace(',', '.')
            return float(price_str) if price_str else None
        except (ValueError, TypeError): return None

    def _clean_stock(self, stock_val: Any) -> str:
        if pd.isna(stock_val): return "не указан"
        stock_str = str(stock_val).lower().strip()
        if any(w in stock_str for w in ['наличи', 'есть', '+']): return "в наличии"
        if any(w in stock_str for w in ['заказ', 'ожид']): return "под заказ"
        numbers = re.findall(r'\d+', stock_str)
        return numbers[0] if numbers else "не указан"

    def get_cascade_summary(self, result: Dict) -> str:
        summary_lines = ["--- Сводка обработки прайс-листа ---"]
        summary_lines.append(f"Статус: {'Успешно' if result.get('success') else 'Ошибка'}")
        summary_lines.append(f"Финальный метод: {result.get('final_method', 'N/A')}")
        summary_lines.append(f"Количество извлеченных товаров: {len(result.get('products', []))}")
        if not result.get('success'):
            summary_lines.append(f"Сообщение об ошибке: {result.get('error', 'N/A')}")
        summary_lines.append("\n--- Детальный лог ---")
        summary_lines.extend(result.get('cascade_log', ["Лог отсутствует."]))
        return "\n".join(summary_lines)

    def _process_with_heuristics(self, file_path: str) -> Dict:
        log = []
        try:
            # Проверяем, есть ли несколько листов в Excel
            file_ext = os.path.splitext(file_path)[1].lower()
            all_products = []
            
            if file_ext in ['.xlsx', '.xls']:
                excel_file = pd.ExcelFile(file_path)
                sheet_names = excel_file.sheet_names
                log.append(f"Файл содержит {len(sheet_names)} лист(ов): {sheet_names}")
                
                # Обрабатываем каждый лист
                for sheet_name in sheet_names:
                    log.append(f"\n--- Обработка листа '{sheet_name}' ---")
                    try:
                        df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
                        log.append(f"Лист '{sheet_name}' прочитан: {df.shape[0]} строк, {df.shape[1]} колонок")
                        
                        # Обрабатываем лист теми же методами
                        products = self._process_single_sheet(df, sheet_name, log)
                        all_products.extend(products)
                        
                    except Exception as sheet_error:
                        log.append(f"Ошибка обработки листа '{sheet_name}': {sheet_error}")
                        continue
                
                if all_products:
                    log.append(f"\nВсего извлечено товаров со всех листов: {len(all_products)}")
                    return {"success": True, "products": all_products, "log": log}
                else:
                    return {"success": False, "products": [], "log": log, "error": "Не найдено товаров ни на одном листе"}
            else:
                # Для других форматов читаем как обычно
                df = pd.read_excel(file_path, header=None)
                log.append(f"Файл {file_path} успешно прочитан в DataFrame.")
                products = self._process_single_sheet(df, "main", log)
                return {"success": True, "products": products, "log": log} if products else {"success": False, "products": [], "log": log, "error": "Не найдено товаров"}
                
        except Exception as e:
            log.append(f"Не удалось прочитать файл: {e}")
            return {"success": False, "products": [], "log": log, "error": f"Ошибка чтения: {e}"}
    
    def _process_single_sheet(self, df: pd.DataFrame, sheet_name: str, log: List[str]) -> List[Dict]:
        """Обработка одного листа Excel"""
        header_row_index, header = self._find_header_row(df, log)
        if header_row_index is None:
            log.append(f"Лист '{sheet_name}': Не удалось найти строку заголовка.")
            return []  # Возвращаем пустой список, а не словарь
        
        df.columns = [f"col_{i}" for i in range(len(df.columns))]
        header_map = {f"col_{i}": str(h) for i, h in enumerate(header)}
        
        data_df = df.iloc[header_row_index + 1:].reset_index(drop=True)
        log.append(f"Лист '{sheet_name}': Данные для обработки подготовлены, начиная со строки {header_row_index + 1}.")

        name_col, price_col, stock_col = self._map_columns(header_map, data_df, log)

        if not name_col or not price_col:
            log.append(f"Лист '{sheet_name}': Не удалось определить обязательные колонки 'Наименование' или 'Цена'.")
            return []  # Возвращаем пустой список

        products = self._extract_products_with_subheaders(data_df, name_col, price_col, stock_col, log, header_map)

        if not products:
            log.append(f"Лист '{sheet_name}': Не удалось извлечь ни одного товара.")
            return []
        
        log.append(f"Лист '{sheet_name}': Извлечено {len(products)} товаров.")
        return products

    def _find_header_row(self, df: pd.DataFrame, log: List[str]) -> Tuple[Optional[int], Optional[List[str]]]:
        header_keywords = ['наимен', 'товар', 'цена', 'кол-во', 'остат', 'артикул', 'руб', 'гост', 'н-ра', 'описание']
        best_row_index = -1
        max_matches = 0

        for i, row in df.head(20).iterrows():
            row_str = ' '.join(str(x) for x in row.dropna() if x).lower()
            if not row_str: continue
            
            matches = sum(1 for kw in header_keywords if kw in row_str)
            
            if matches > 1 and matches > max_matches:
                max_matches = matches
                best_row_index = i
        
        if best_row_index != -1:
            log.append(f"Найдена строка заголовка (индекс {best_row_index}) с {max_matches} совпадениями.")
            return best_row_index, list(df.iloc[best_row_index])
        
        log.append("Строка заголовка не найдена, используется первая строка.")
        return 0, list(df.iloc[0])

    def _map_columns(self, header_map: Dict, df: pd.DataFrame, log: List[str]) -> Tuple[Optional[str], Optional[str], Optional[str]]:
        price_kw = ['цена', 'price', 'стоим', 'cost', 'value', 'руб', 'rub', 'сумма']
        stock_kw = ['остат', 'кол-во', 'налич', 'stock', 'qty', 'amount', 'balance', 'количество', 'склад']
        name_kw = ['наимен', 'товар', 'product', 'item', 'описан', 'nomenkl', 'назв', 'продукт']

        def find_header(keywords):
            for col_idx, header_name in header_map.items():
                header_str = str(header_name).lower()
                if header_str != 'nan' and any(k in header_str for k in keywords):
                    return col_idx
            return None

        # Ищем колонки по заголовкам
        name_col = find_header(name_kw)
        price_col = find_header(price_kw)
        stock_col = find_header(stock_kw)
        
        # Если не нашли колонку с названием, ищем по содержимому
        if not name_col:
            log.append("Не найдена колонка названий по заголовкам. Пробуем определить по содержимому...")
            for col_idx in header_map.keys():
                text_like_count = 0
                for i in range(min(10, len(df))):
                    val = str(df.iloc[i].get(col_idx, '')).strip()
                    # Проверяем, похоже ли на текст (не число и не пустое)
                    if val and val != 'nan' and re.search(r'[а-яА-Яa-zA-Z]', val) and not re.match(r'^\d+[\.,\d]*$', val):
                        text_like_count += 1
                
                if text_like_count >= 5:  # Если больше половины значений похожи на текст
                    name_col = col_idx
                    log.append(f"Вероятная колонка названий найдена по содержимому: {col_idx}")
                    break
        
        # Если не нашли колонку с ценой по ключевым словам, попробуем найти по данным
        if not price_col:
            log.append("Не найдена колонка цены по заголовкам. Пробуем определить по содержимому...")
            for col_idx in header_map.keys():
                if col_idx == name_col:  # Пропускаем колонку с названием
                    continue
                    
                price_like_count = 0
                for i in range(min(10, len(df))):
                    val = str(df.iloc[i].get(col_idx, '')).strip()
                    # Проверяем, похоже ли на цену (только цифры, точки, запятые)
                    if val and val != 'nan' and re.match(r'^\d+[\.,\d\s]*$', val):
                        price_like_count += 1
                
                if price_like_count >= 5:  # Если больше половины значений похожи на цены
                    price_col = col_idx
                    log.append(f"Вероятная колонка цены найдена по содержимому: {col_idx}")
                    break
        
        # Если не нашли колонку с остатком, ищем по содержимому
        if not stock_col:
            log.append("Не найдена колонка остатков по заголовкам. Пробуем определить по содержимому...")
            for col_idx in header_map.keys():
                if col_idx in [name_col, price_col]:  # Пропускаем уже найденные колонки
                    continue
                    
                stock_like_count = 0
                for i in range(min(10, len(df))):
                    val = str(df.iloc[i].get(col_idx, '')).strip().lower()
                    # Проверяем, похоже ли на остаток
                    if val and val != 'nan' and (re.match(r'^\d+$', val) or any(w in val for w in ['наличи', 'заказ', 'есть', 'нет'])):
                        stock_like_count += 1
                
                if stock_like_count >= 3:  # Более мягкий критерий для остатков
                    stock_col = col_idx
                    log.append(f"Вероятная колонка остатков найдена по содержимому: {col_idx}")
                    break
        
        # Если все еще не нашли название, берем первую колонку по умолчанию
        if not name_col:
            name_col = 'col_0'
            log.append("Используем первую колонку как колонку названий по умолчанию")
        
        log.append(f"Финальный маппинг: Name='{header_map.get(name_col)}' ({name_col}), Price='{header_map.get(price_col)}' ({price_col}), Stock='{header_map.get(stock_col)}' ({stock_col})")
        
        return name_col, price_col, stock_col

    def _extract_products_with_subheaders(self, df: pd.DataFrame, name_col: str, price_col: str, stock_col: Optional[str], log: List[str], header_map: Dict) -> List[Dict]:
        products = []
        current_subheader = ""
        for index, row in df.iterrows():
            name = str(row.get(name_col, "")).strip()
            
            # Пропускаем строки с "nan" или пустыми названиями
            if not name or name.lower() in ['nan', 'none', '']:
                continue
                
            is_subheader = name and all(pd.isna(v) or str(v).strip() == "" or str(v).strip().lower() == 'nan' for k, v in row.items() if k != name_col)
            
            if is_subheader:
                current_subheader = name
                log.append(f"Обнаружен подзаголовок: '{current_subheader}'")
                continue

            full_name = f"{current_subheader} {name}".strip()
            
            price = 0
            price_raw = str(row.get(price_col, "")).strip()
            if price_raw and price_raw not in ['nan', 'None', '-', '']:
                # Логируем исходное значение цены для отладки
                logger.debug(f"Обрабатываем цену: '{price_raw}'")
                
                # Удаляем все пробельные символы включая неразрывные пробелы
                price_raw = price_raw.replace('\xa0', ' ').replace('\u00a0', ' ')
                
                # Если цена содержит дробь через пробел (например "1 234,56")
                # или через точку как разделитель тысяч (например "1.234,56")
                price_raw = price_raw.replace(' ', '')
                
                # Заменяем запятую на точку для десятичных
                if ',' in price_raw and '.' in price_raw:
                    # Если есть и точка и запятая, предполагаем что точка - разделитель тысяч
                    price_raw = price_raw.replace('.', '').replace(',', '.')
                else:
                    price_raw = price_raw.replace(',', '.')
                
                # Извлекаем только числа и точку
                clean_price = re.sub(r'[^\d\.]', '', price_raw)
                
                if clean_price:
                    try:
                        price = float(clean_price)
                        logger.debug(f"Успешно извлечена цена: {price}")
                    except (ValueError, TypeError) as e:
                        logger.debug(f"Не удалось преобразовать цену '{clean_price}': {e}")
            
            stock = 100
            if stock_col and pd.notna(row.get(stock_col)):
                stock_raw = str(row[stock_col]).lower().strip()
                if stock_raw != 'nan':
                    if any(w in stock_raw for w in ['нет', '0', 'под заказ', 'ожид', 'отсут']): 
                        stock = 0
                    elif any(w in stock_raw for w in ['есть', 'в наличии', 'налич', 'много']): 
                        stock = 100
                    else:
                        stock_numbers = re.findall(r'\d+', stock_raw)
                        if stock_numbers: 
                            stock = int(stock_numbers[0])

            # Добавляем товар только если есть валидное название
            if full_name and full_name.strip() and full_name.lower() != 'nan':
                products.append({"name": full_name, "price": price, "stock": stock})
                if price == 0:
                    logger.warning(f"Товар с нулевой ценой: {full_name[:50]}...")
        
        log.append(f"Извлечено {len(products)} товаров.")
        return products
    
    def _read_file_safely(self, file_path: str) -> Optional[pd.DataFrame]:
        """Безопасное чтение файла"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.csv':
                # Для CSV пробуем разные кодировки
                for encoding in ['utf-8', 'cp1251', 'latin1']:
                    try:
                        return pd.read_csv(file_path, encoding=encoding, dtype=str)
                    except:
                        continue
                        
            elif file_ext in ['.xlsx', '.xls']:
                # Для Excel пробуем разные движки
                if file_ext == '.xls':
                    try:
                        return pd.read_excel(file_path, header=None, dtype=str, engine='xlrd')
                    except:
                        try:
                            return pd.read_excel(file_path, header=None, dtype=str, engine='openpyxl')
                        except:
                            return pd.read_excel(file_path, header=None, dtype=str)
                else:
                    return pd.read_excel(file_path, header=None, dtype=str)
                    
            return None
            
        except Exception as e:
            logger.error(f"Ошибка чтения файла {file_path}: {e}")
            return None
    
    def _select_best_overall_result(self, level1_result: Dict, level2_result: Dict) -> Optional[Dict]:
        """Выбор лучшего результата из всех уровней"""
        candidates = []
        
        if level1_result.get('success') and level1_result.get('products'):
            candidates.append({
                'method': 'level1_smart',
                'products': level1_result['products'],
                'count': len(level1_result['products'])
            })
            
        if level2_result.get('success') and level2_result.get('products'):
            candidates.append({
                'method': 'level2_bruteforce', 
                'products': level2_result['products'],
                'count': len(level2_result['products'])
            })
        
        if not candidates:
            return None
            
        # Выбираем кандидата с наибольшим количеством товаров
        best = max(candidates, key=lambda x: x['count'])
        return best
        
    # Этот метод закомментирован, так как он использует несуществующие level1_processor и level2_processor
    # def process_file_cascade(self, file_path: str, file_name: str = "") -> Dict:
    #     """
    #     Каскадная обработка файла с тремя уровнями сложности
    #     """
    #     logger.info(f"🎯 КАСКАДНАЯ ОБРАБОТКА: {file_name}")
    #     
    #     result = {
    #         'success': False,
    #         'final_method': None,
    #         'products': [],
    #         'cascade_log': [],
    #         'all_attempts': {}
    #     }
    #     
    #     try:
    #         # Читаем файл для передачи в процессоры
    #         df = self._read_file_safely(file_path)
    #         if df is None:
    #             result['cascade_log'].append("❌ Ошибка чтения файла")
    #             return result
    #         
    #         # УРОВЕНЬ 1: УМНАЯ СИСТЕМА
    #         logger.info("🔥 УРОВЕНЬ 1: Умная система")
    #         level1_result = self.level1_processor.process_price_list(df, file_name)
    #         result['all_attempts']['level1'] = level1_result
    #         
    #         if level1_result['success'] and len(level1_result['products']) >= 5:
    #             # Успех на уровне 1!
    #             result['success'] = True
    #             result['final_method'] = 'level1_smart'
    #             result['products'] = level1_result['products']
    #             result['cascade_log'].append(f"✅ УРОВЕНЬ 1 УСПЕХ: {len(level1_result['products'])} товаров")
    #             logger.info(f"✅ Уровень 1 успешен: {len(level1_result['products'])} товаров")
    #             return result
    #         else:
    #             result['cascade_log'].append(f"⚠️ УРОВЕНЬ 1: {len(level1_result.get('products', []))} товаров (недостаточно)")
    #             
    #         # УРОВЕНЬ 2: БРУТФОРС АНАЛИЗ
    #         logger.info("🔥 УРОВЕНЬ 2: Брутфорс анализ")
    #         level2_result = self.level2_processor.process_complex_file(file_path, file_name)
    #         result['all_attempts']['level2'] = level2_result
    #         
    #         if level2_result['success'] and len(level2_result['products']) > 0:
    #             # Сравниваем результаты уровня 1 и 2
    #             level1_count = len(level1_result.get('products', []))
    #             level2_count = len(level2_result['products'])
    #             
    #             if level2_count > level1_count:
    #                 # Уровень 2 лучше
    #                 result['success'] = True
    #                 result['final_method'] = 'level2_bruteforce'
    #                 result['products'] = level2_result['products']
    #                 result['cascade_log'].append(f"✅ УРОВЕНЬ 2 ЛУЧШЕ: {level2_count} товаров (vs {level1_count})")
    #                 logger.info(f"✅ Уровень 2 лучше: {level2_count} vs {level1_count} товаров")
    #                 return result
    #             elif level1_count > 0:
    #                 # Уровень 1 все же лучше
    #                 result['success'] = True  
    #                 result['final_method'] = 'level1_smart'
    #                 result['products'] = level1_result['products']
    #                 result['cascade_log'].append(f"✅ УРОВЕНЬ 1 ЛУЧШЕ: {level1_count} товаров (vs {level2_count})")
    #                 return result
    #         else:
    #             result['cascade_log'].append(f"⚠️ УРОВЕНЬ 2: {len(level2_result.get('products', []))} товаров")
    #             
    #         # УРОВЕНЬ 3: ПОЛНЫЙ LLM АНАЛИЗ (пока заглушка)
    #         logger.info("🔥 УРОВЕНЬ 3: Полный LLM анализ")
    #         result['cascade_log'].append("⏳ УРОВЕНЬ 3: Требует реализации")
    #         
    #         # Если дошли сюда - берем лучший результат из уровней 1-2
    #         best_result = self._select_best_overall_result(level1_result, level2_result)
    #         if best_result:
    #             result['success'] = True
    #             result['final_method'] = best_result['method']
    #             result['products'] = best_result['products']
    #             result['cascade_log'].append(f"✅ ЛУЧШИЙ РЕЗУЛЬТАТ: {best_result['method']} - {len(best_result['products'])} товаров")
    #         else:
    #             result['cascade_log'].append("❌ ВСЕ УРОВНИ НЕУСПЕШНЫ")
    #             
    #     except Exception as e:
    #         logger.error(f"Ошибка каскадной обработки: {e}")
    #         result['cascade_log'].append(f"❌ КРИТИЧЕСКАЯ ОШИБКА: {str(e)}")
    #         
    #     return result 